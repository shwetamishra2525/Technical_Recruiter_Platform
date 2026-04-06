from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from typing import List
from models import UserModel, JobDescriptionModel, CandidateProfileModel, InterviewModel
from database import users_collection, jds_collection, candidates_collection, interviews_collection
from auth import get_password_hash, verify_password, create_access_token, get_current_user
from ai_service import gemini_service
from services.mail_service import send_status_email
from bson import ObjectId
from difflib import SequenceMatcher
import fitz # PyMuPDF

router = APIRouter()

# --- Auth Routes ---
@router.post("/auth/signup", status_code=status.HTTP_201_CREATED)
async def signup(user: UserModel):
    try:
        existing_user = await users_collection.find_one({"email": user.email})
        if existing_user:
            raise HTTPException(status_code=400, detail="Email already registered")
        
        user.password = get_password_hash(user.password)
        new_user = await users_collection.insert_one(user.model_dump(by_alias=True, exclude=["id"]))
        created_user = await users_collection.find_one({"_id": new_user.inserted_id})
        return {"id": str(created_user["_id"]), "email": created_user["email"], "role": created_user["role"]}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/auth/login")
async def login(login_data: dict):
    user = await users_collection.find_one({"email": login_data.get("email")})
    if not user or not verify_password(login_data.get("password"), user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    access_token = create_access_token(data={"sub": user["email"], "role": user["role"], "id": str(user["_id"])})
    return {"access_token": access_token, "token_type": "bearer", "role": user["role"]}

# --- User Profile Routes ---
@router.get("/user/profile-status")
async def get_user_profile_status(current_user: dict = Depends(get_current_user)):
    user = await users_collection.find_one({"_id": ObjectId(current_user["id"])})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
        
    if current_user["role"] == "hr":
        completion = 20
        missing = []
        if user.get("company_name"):
            completion += 40
        else:
            missing.append("company_name")
            
        if user.get("phone"):
            completion += 40
        else:
            missing.append("phone")
            
        return {
            "role": "hr", 
            "completion_percentage": completion, 
            "missing_fields": missing, 
            "user_details": {
                "company_name": user.get("company_name", ""), 
                "phone": user.get("phone", ""),
                "photo_url": user.get("photo_url", "")
            }
        }
    else:
        # Candidate
        completion = 20
        candidate = await candidates_collection.find_one({"user_id": current_user["id"]})
        if candidate and candidate.get("resume_text"):
            completion += 80
            
        return {
            "role": "candidate", 
            "completion_percentage": completion,
            "user_details": {
                "photo_url": user.get("photo_url", "")
            }
        }

@router.post("/user/update-profile")
async def update_user_profile(profile_data: dict, current_user: dict = Depends(get_current_user)):
    update_fields = {}
    for field in ["company_name", "phone", "gender", "location", "dob"]:
        if field in profile_data:
            update_fields[field] = profile_data[field]
        
    if update_fields:
        await users_collection.update_one(
            {"_id": ObjectId(current_user["id"])},
            {"$set": update_fields}
        )
    return {"message": "Profile updated successfully"}

@router.post("/user/upload-photo")
async def upload_user_photo(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    import shutil
    import os
    UPLOAD_DIR = "/tmp/uploads"
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    file_extension = file.filename.split(".")[-1]
    filename = f"{current_user['id']}_profile.{file_extension}"
    file_location = f"{UPLOAD_DIR}/{filename}"
    
    with open(file_location, "wb+") as file_object:
        shutil.copyfileobj(file.file, file_object)
        
    photo_url = f"/uploads/{filename}"
    await users_collection.update_one(
        {"_id": ObjectId(current_user["id"])},
        {"$set": {"photo_url": photo_url}}
    )
    return {"photo_url": photo_url}

@router.get("/user/profile")
async def get_full_user_profile(current_user: dict = Depends(get_current_user)):
    user = await users_collection.find_one({"_id": ObjectId(current_user["id"])})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
        
    # Calculate completion percentage
    fields_to_check = ["phone", "gender", "location", "dob", "photo_url"]
    if current_user["role"] == "hr":
        fields_to_check.append("company_name")
        
    filled_fields = sum(1 for field in fields_to_check if user.get(field))
    
    # Base 20% + 80% distributed among fields
    field_weight = 80 / len(fields_to_check)
    completion = 20 + (filled_fields * field_weight)
    
    # Fetch History
    history = []
    import pymongo
    if current_user["role"] == "candidate":
        candidate = await candidates_collection.find_one({"user_id": current_user["id"]})
        if candidate:
            interviews = interviews_collection.find({"candidate_id": str(candidate["_id"])}).sort("created_at", pymongo.DESCENDING)
            async for iv in interviews:
                jd = await jds_collection.find_one({"_id": ObjectId(iv["jd_id"])})
                history.append({
                    "id": str(iv["_id"]),
                    "title": jd["title"] if jd else "Unknown Role",
                    "status": iv.get("status", "pending"),
                    "score": round(iv.get("total_score", 0), 1),
                    "date": iv.get("created_at")
                })
    else:
        # HR history
        jds = jds_collection.find({"hr_id": current_user["id"]}).sort("created_at", pymongo.DESCENDING)
        async for jd in jds:
            applicant_count = await interviews_collection.count_documents({"jd_id": str(jd["_id"])})
            history.append({
                "id": str(jd["_id"]),
                "title": jd["title"],
                "applicants": applicant_count,
                "date": jd.get("created_at")
            })
            
    return {
        "user": {
            "email": user.get("email"),
            "role": user.get("role"),
            "phone": user.get("phone", ""),
            "company_name": user.get("company_name", ""),
            "photo_url": user.get("photo_url", ""),
            "gender": user.get("gender", ""),
            "location": user.get("location", ""),
            "dob": user.get("dob", ""),
            "created_at": user.get("created_at")
        },
        "completion_percentage": min(100, round(completion)),
        "history": history
    }

# --- HR Routes ---
@router.post("/hr/create-jd", response_model=JobDescriptionModel)
async def create_jd(jd: JobDescriptionModel, current_user: dict = Depends(get_current_user)):
    try:
        print(f"DEBUG: create_jd called by {current_user['email']} (Role: {current_user['role']})")
        print(f"DEBUG: JD Payload: {jd.model_dump()}")
        
        if current_user["role"] != "hr":
            raise HTTPException(status_code=403, detail="Only HR can create JDs")
        
        jd.hr_id = current_user["id"]
        new_jd = await jds_collection.insert_one(jd.model_dump(by_alias=True, exclude=["id"]))
        created_jd = await jds_collection.find_one({"_id": new_jd.inserted_id})
        print(f"DEBUG: JD Created: {created_jd['_id']}")
        return created_jd
    except Exception as e:
        import traceback
        traceback.print_exc()
        # If it's already an HTTPException, re-raise it
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@router.post("/hr/upload-jd", response_model=JobDescriptionModel)
async def upload_jd(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    if current_user["role"] != "hr":
        raise HTTPException(status_code=403, detail="Only HR can upload JDs")
    
    content = await file.read()
    
    text_content = ""
    try:
        filename_lower = file.filename.lower()
        if filename_lower.endswith(".pdf"):
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc:
                    text_content += page.get_text()
        elif filename_lower.endswith(".docx"):
            import docx
            import io
            document = docx.Document(io.BytesIO(content))
            text_content = "\n".join([para.text for para in document.paragraphs])
        else:
            text_content = content.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Extraction Error: {e}")
        text_content = content.decode("utf-8", errors="ignore")

    if len(text_content.strip()) < 20:
         raise HTTPException(status_code=400, detail="Document content is too short or unreadable.")

    print("DEBUG EXTRACTED JD TEXT:", repr(text_content[:500]))
    parsed_data = await gemini_service.parse_jd(text_content)
    
    if parsed_data.get("title") == "Invalid Document" or parsed_data.get("title") == "Unknown":
        error_msg = parsed_data.get("error", "The uploaded document does not appear to be a valid Job Description.")
        raise HTTPException(status_code=400, detail=error_msg)

    jd_model = JobDescriptionModel(
        title=parsed_data.get("title", "Unknown Title"),
        skills=parsed_data.get("skills", []),
        requirements=parsed_data.get("requirements", text_content[:500]),
        hr_id=current_user["id"]
    )
    
    new_jd = await jds_collection.insert_one(jd_model.model_dump(by_alias=True, exclude=["id"]))
    created_jd = await jds_collection.find_one({"_id": new_jd.inserted_id})
    return created_jd

@router.get("/hr/dashboard")
async def get_dashboard_stats(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "hr":
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    import pymongo
    
    # 1. Fetch JDs created by this HR first
    jds_cursor = jds_collection.find({"hr_id": current_user["id"]}).sort("created_at", pymongo.DESCENDING)
    jds_list = []
    hr_jd_ids = []
    
    async for jd in jds_cursor:
        hr_jd_ids.append(str(jd["_id"]))
        jds_list.append({
            "id": str(jd["_id"]),
            "title": jd["title"],
            "skills": jd["skills"],
            "requirements": jd["requirements"],
            "created_at": jd.get("created_at")
        })
    
    # 2. Filter query for interviews
    interview_query = {"jd_id": {"$in": hr_jd_ids}}

    # (We calculate KPIs inside the loop below to ensure deduplication matches the GUI)

    # 4. Interviews (Listing ONLY applications for this HR's JDs)
    interviews_cursor = interviews_collection.find(interview_query).sort("created_at", pymongo.DESCENDING).limit(500)
    applications_list = []
    
    unique_total = 0
    unique_ready = 0
    unique_finished = 0
    
    async for interview in interviews_cursor:
        candidate_id = str(interview.get("candidate_id", ""))
        jd_id = str(interview.get("jd_id", ""))
        # Removed deduplication to ensure all applications/attempts are visible to HR as requested
        
        # Calculate KPIs
        unique_total += 1
        status = interview.get("status", "New")
        if status == "completed":
            unique_ready += 1
        elif status in ["Approved", "Reject"]:
            unique_finished += 1
        
        try:
            cand_id_obj = ObjectId(interview["candidate_id"])
        except Exception:
            cand_id_obj = None
            
        try:
            jd_id_obj = ObjectId(interview["jd_id"])
        except Exception:
            jd_id_obj = None

        candidate = await candidates_collection.find_one({"_id": cand_id_obj}) if cand_id_obj else None
        jd = await jds_collection.find_one({"_id": jd_id_obj}) if jd_id_obj else None
        
        # Use the name snapped at the time of the test if multiple users share an account
        candidate_name = interview.get("candidate_name")
        if not candidate_name and candidate:
            candidate_name = candidate.get("full_name", "Unknown Applicant")
        elif not candidate_name:
            candidate_name = "Deleted Candidate"
            
        applications_list.append({
            "id": str(interview["_id"]), # Navigation now happens via interview_id
            "candidate_id": str(candidate["_id"]) if candidate else None,
            "name": candidate_name,
            "role": jd.get("title") if jd else "General Applicant",
            "score": interview.get("total_score", 0),
            "status": status,
            "created_at": interview.get("created_at")
        })
        
    completion_rate = round((unique_finished / unique_total * 100), 1) if unique_total > 0 else 0
        
    return {
        "stats": {
            "total_applicants": unique_total, # Renamed logically in response but frontend will map it
            "pending_reviews": unique_ready,
            "completion_rate": f"{completion_rate}% ({unique_finished}/{unique_total})"
        },
        "candidates": applications_list, # Kept key 'candidates' for frontend compatibility but content is interviews
        "jds": jds_list
    }

@router.delete("/hr/job/{jd_id}")
async def delete_jd(jd_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "hr":
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    result = await jds_collection.delete_one({"_id": ObjectId(jd_id)})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Job description not found")
        
    return {"message": "Job description deleted successfully"}

@router.post("/hr/update-status/{interview_id}")
async def update_interview_status(interview_id: str, status_data: dict, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "hr":
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    new_status = status_data.get("status")
    interview = await interviews_collection.find_one({"_id": ObjectId(interview_id)})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview session not found")
        
    candidate = await candidates_collection.find_one({"_id": ObjectId(interview["candidate_id"])})
    if not candidate:
         raise HTTPException(status_code=404, detail="Candidate record not found")

    # Update this specific session
    await interviews_collection.update_one(
        {"_id": ObjectId(interview_id)},
        {"$set": {"status": new_status}}
    )

    # Send Email to Candidate
    user = await users_collection.find_one({"_id": ObjectId(candidate["user_id"])})
    if user:
        await send_status_email(user["email"], new_status, candidate["full_name"])
    
    return {"message": f"Application {new_status} and email sent."}

@router.get("/hr/interview/{interview_id}")
async def get_interview_details(interview_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "hr":
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    interview = await interviews_collection.find_one({"_id": ObjectId(interview_id)})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview session not found")
        
    candidate = await candidates_collection.find_one({"_id": ObjectId(interview["candidate_id"])})
    if not candidate:
         raise HTTPException(status_code=404, detail="Candidate record not found")
    
    return {
        "profile": {
            "full_name": candidate.get("full_name"),
            "email": candidate.get("email"),
            "skills": candidate.get("skills"),
            "experience_years": candidate.get("experience_years"),
            "resume_text": candidate.get("resume_text")[:500] + "..." if candidate.get("resume_text") else ""
        },
        "interview": {
            "id": str(interview["_id"]),
            "status": interview.get("status") if interview else "Not Started",
            "total_score": interview.get("total_score") if interview else 0,
            "questions": interview.get("questions", []) if interview else []
        }
    }

@router.get("/candidate/profile")
async def get_candidate_profile(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "candidate":
        raise HTTPException(status_code=403, detail="Only candidates can access this")
    
    candidate = await candidates_collection.find_one({"user_id": current_user["id"]})
    if not candidate:
        return {"profile": None, "application_status": None}
    
    # Check status
    import pymongo
    latest_interview = await interviews_collection.find_one(
        {"candidate_id": str(candidate["_id"])},
        sort=[("created_at", pymongo.DESCENDING)]
    )
    
    return {
        "profile": {
            "full_name": candidate.get("full_name"),
            "skills": candidate.get("skills"),
            "experience_years": candidate.get("experience_years")
        },
        "application_status": latest_interview.get("status") if latest_interview else None
    }

# --- Candidate Routes ---



@router.post("/candidate/upload-resume")
async def upload_resume(
    file: UploadFile = File(...), 
    current_user: dict = Depends(get_current_user)
):
    if current_user["role"] != "candidate":
        raise HTTPException(status_code=403, detail="Only candidates can upload resumes")
    
    # Read file content
    content = await file.read()
    
    # Extract text using PyMuPDF (fitz) or docx
    text_content = ""
    try:
        filename_lower = file.filename.lower()
        if filename_lower.endswith(".pdf"):
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc:
                    text_content += page.get_text()
        elif filename_lower.endswith(".docx"):
            import docx
            import io
            document = docx.Document(io.BytesIO(content))
            text_content = "\n".join([para.text for para in document.paragraphs])
        else:
            text_content = content.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Extraction Error: {e}")
        text_content = content.decode("utf-8", errors="ignore") # Fallback

    print(f"DEBUG: Extracted text length: {len(text_content)}")
    if len(text_content.strip()) < 50:
         raise HTTPException(status_code=400, detail="Document content is too short or unreadable. Please upload a text-based PDF (not scanned image).")

    # Parse with AI
    parsed_data = await gemini_service.parse_resume(text_content)
    print(f"DEBUG: AI Raw Output: {parsed_data}")
    
    full_name_lower = str(parsed_data.get("full_name", "")).lower()
    
    # Catch any variation of the error message winding up in the full_name field due to AI hallucination
    if (parsed_data.get("valid") is False or 
        "invalid document" in full_name_lower or 
        "unknown" in full_name_lower or 
        "does not appear" in full_name_lower or 
        "not a resume" in full_name_lower):
        
        error_msg = parsed_data.get("error")
        if not error_msg or len(error_msg) < 5:
            error_msg = "The uploaded document does not appear to be a valid resume or AI processing failed."
            
        raise HTTPException(status_code=400, detail=error_msg)

    candidate_profile = CandidateProfileModel(
        user_id=current_user["id"],
        full_name=parsed_data.get("full_name", "Unknown"),
        resume_text=text_content,
        skills=parsed_data.get("skills", []),
        experience_years=parsed_data.get("experience_years", 0),
        matched_jds={} # Clear the cache since they updated their resume!
    )
    
    existing = await candidates_collection.find_one({"user_id": current_user["id"]})
    if existing:
        await candidates_collection.update_one({"_id": existing["_id"]}, {"$set": candidate_profile.model_dump(exclude=["id"])})
    else:
        await candidates_collection.insert_one(candidate_profile.model_dump(by_alias=True, exclude=["id"]))
        
    return {"message": "Resume uploaded successfully", "data": parsed_data}

@router.get("/candidate/jobs")
async def get_all_jobs(current_user: dict = Depends(get_current_user)):
    # 1. Get all JDs and deduplicate by title using fuzzy matching
    jds_cursor = jds_collection.find()
    jobs = []
    seen_titles = [] # Store lowercased titles
    
    async for jd in jds_cursor:
        lower_title = jd["title"].lower().strip()
        is_duplicate = False
        
        for seen in seen_titles:
            # Check similarity ratio (e.g. "react developer" vs "react devloepr")
            similarity = SequenceMatcher(None, lower_title, seen).ratio()
            if similarity > 0.95:
                is_duplicate = True
                break
                
        if not is_duplicate:
            seen_titles.append(lower_title)
            jobs.append({
                "id": str(jd["_id"]),
                "title": jd["title"],
                "skills": jd["skills"],
                "requirements": jd["requirements"],
                "match_score": 0, # Default
                "match_reason": ""
            })
    
    # 2. Get Candidate Profile (Resume)
    if current_user["role"] == "candidate":
        candidate = await candidates_collection.find_one({"user_id": current_user["id"]})
        if candidate and candidate.get("resume_text") and len(candidate["resume_text"]) > 50:
            print(f"DEBUG: Checking cache for candidate {current_user['email']}")
            
            matched_jds_cache = candidate.get("matched_jds", {})
            unmatched_jobs = [j for j in jobs if j["id"] not in matched_jds_cache]
            
            if unmatched_jobs:
                print(f"DEBUG: Found {len(unmatched_jobs)} new JDs. Calling AI matcher...")
                # 3. Call AI to Match (ONLY for missing ones!)
                matches = await gemini_service.match_jds(candidate["resume_text"], unmatched_jobs)
                
                for m in matches:
                    matched_jds_cache[m["id"]] = m
                    
                # Store the updated cache back to DB
                await candidates_collection.update_one(
                    {"_id": candidate["_id"]},
                    {"$set": {"matched_jds": matched_jds_cache}}
                )
            
            # 4. Merge Scores
            for job in jobs:
                job_id = job["id"]
                if job_id in matched_jds_cache:
                    job["match_score"] = matched_jds_cache[job_id].get("match_score", 0)
                    job["match_reason"] = matched_jds_cache[job_id].get("reason", "")
            
            # 5. Sort by Score (Descending). Since Python sort is stable, newer JDs naturally bubble up on ties.
            jobs.sort(key=lambda x: x.get("match_score", 0), reverse=True)
            
            # 6. Top 3 results as requested
            jobs = jobs[:3]
            
    return jobs

@router.post("/candidate/start-interview/{jd_id}")
async def start_interview(jd_id: str, current_user: dict = Depends(get_current_user)):
    candidate = await candidates_collection.find_one({"_id": current_user["id"]}) # wait it's user_id or _id, let's fix the query to find by user_id
    if not candidate:
        candidate = await candidates_collection.find_one({"user_id": current_user["id"]})
        
    if not candidate:
        raise HTTPException(status_code=400, detail="Please upload resume first")
        
    jd = await jds_collection.find_one({"_id": ObjectId(jd_id)})
    if not jd:
        raise HTTPException(status_code=404, detail="Job not found")
        
    # Check if they ALREADY have an interview for this exact JD
    import pymongo
    existing_interview = await interviews_collection.find_one(
        {"candidate_id": str(candidate["_id"]), "jd_id": jd_id},
        sort=[("created_at", pymongo.DESCENDING)]
    )
    
    if existing_interview:
        status = existing_interview.get("status", "").lower()
        if status in ["approved", "completed"]:
            raise HTTPException(status_code=400, detail="You have already completed the application for this role.")
        elif status == "reject":
            # Candidate was rejected previously, simply ignore the old record and let them take a fresh new test
            pass
        else:
            # It's pending, let them resume it instead of creating a duplicate
            return {"interview_id": str(existing_interview["_id"]), "questions": existing_interview.get("questions", [])}
        
    # Generate Questions
    questions = await gemini_service.generate_interview_questions(candidate["resume_text"], jd["requirements"])
    
    interview_data = InterviewModel(
        candidate_id=str(candidate["_id"]),
        candidate_name=candidate.get("full_name", "Unknown Applicant"),
        jd_id=jd_id,
        questions=[{"question": q, "answer": "", "score": 0} for q in questions],
        status="pending"
    )
    
    new_interview = await interviews_collection.insert_one(interview_data.model_dump(by_alias=True, exclude=["id"]))
    return {"interview_id": str(new_interview.inserted_id), "questions": questions}

@router.post("/candidate/submit-answer/{interview_id}")
async def submit_answer(interview_id: str, answer_data: dict):
    print(f"DEBUG: Submitting Answer for {interview_id}. Data: {answer_data}")
    interview = await interviews_collection.find_one({"_id": ObjectId(interview_id)})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
        
    idx = answer_data["question_index"]
    ans = answer_data["answer"]
    
    # Evaluate
    question = interview["questions"][idx]["question"]
    eval_result = await gemini_service.evaluate_answer(question, ans)
    print(f"DEBUG: Eval Result: {eval_result}")
    
    # Update DB
    update_res = await interviews_collection.update_one(
        {"_id": ObjectId(interview_id)},
        {"$set": {
            f"questions.{idx}.answer": ans,
            f"questions.{idx}.score": eval_result.get("score", 0),
            f"questions.{idx}.feedback": eval_result.get("feedback", "")
        }}
    )
    print(f"DEBUG: DB Update Modified Count: {update_res.modified_count}")
    
    # Recalculate Total
    interview = await interviews_collection.find_one({"_id": ObjectId(interview_id)})
    total_score = sum(q.get("score", 0) for q in interview["questions"]) / len(interview["questions"])
    await interviews_collection.update_one({"_id": ObjectId(interview_id)}, {"$set": {"total_score": total_score}})
    
    return eval_result

@router.post("/candidate/submit-all-answers/{interview_id}")
async def submit_all_answers(interview_id: str, payload: dict):
    print(f"DEBUG: Submitting all answers for {interview_id}")
    interview = await interviews_collection.find_one({"_id": ObjectId(interview_id)})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
        
    answers = payload.get("answers", [])
    if len(answers) != len(interview["questions"]):
        print(f"WARNING: Provided answers ({len(answers)}) don't match questions ({len(interview['questions'])})")
        
    qa_pairs = []
    for i, q in enumerate(interview["questions"]):
        ans = answers[i] if i < len(answers) else ""
        qa_pairs.append({
            "question": q["question"],
            "answer": ans
        })
        
    # Bulk evaluate
    eval_results = await gemini_service.evaluate_all_answers(qa_pairs)
    print(f"DEBUG: Bulk Eval Results: {eval_results}")
    
    db_updates = {}
    total_score = 0
    
    for i in range(len(interview["questions"])):
        score = 0
        feedback = ""
        # Safely extract from eval_results if it exists
        if i < len(eval_results) and isinstance(eval_results[i], dict):
            score = eval_results[i].get("score", 0)
            feedback = eval_results[i].get("feedback", "")
            
        db_updates[f"questions.{i}.answer"] = qa_pairs[i]["answer"]
        db_updates[f"questions.{i}.score"] = score
        db_updates[f"questions.{i}.feedback"] = feedback
        
        total_score += score
        
    avg_score = total_score / len(interview["questions"]) if interview["questions"] else 0
    db_updates["total_score"] = avg_score
    db_updates["status"] = "completed"
    
    await interviews_collection.update_one(
        {"_id": ObjectId(interview_id)},
        {"$set": db_updates}
    )
    
    return {"message": "All answers submitted and evaluated successfully", "total_score": avg_score}
